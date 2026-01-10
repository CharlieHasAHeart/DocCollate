import json
import logging
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document

# --- 配置日志，方便追踪脚本运行情况 ---
logging.basicConfig(level=logging.WARNING, format='%(asctime)s - %(levelname)s - %(message)s')

# --- 常量定义，方便维护 ---
CHECKED_SYMBOL = "☑"
UNCHECKED_SYMBOL = "☐"
CHECKED_DOT = "●"
UNCHECKED_DOT = "○"


# ===================================================================
#           最终版核心替换函数 (保证在纯表格文档中生效)
# ===================================================================

def docx_replace_text(doc, replacements: dict) -> None:
    """
    遍历Word文档的所有部分（段落和表格），执行一个可靠的文本替换。
    此版本通过重写段落的 .text 属性并直接在文档结构中操作来保证替换生效。
    """
    # 1. 替换文档顶层的独立段落
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # 2. 替换所有表格单元格内的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def replace_in_paragraph(paragraph, replacements: dict) -> None:
    """
    辅助函数，在一个段落内循环替换所有找到的占位符。
    通过重写 .text 属性来确保替换的可靠性。
    此版本经过修改，可以同时处理带软回车 (在python-docx中为 \v ) 的占位符。
    """
    # 在 python-docx 库中，Word的软回车(↵)通常被解析为垂直制表符 '\v'
    soft_break = "\v"

    # 只要段落中还存在 '{{'，就持续进行替换
    while "{{" in paragraph.text:
        original_text = paragraph.text
        for key, value in replacements.items():
            # 使用 str() 确保所有值（包括None）都能被处理
            value_str = str(value if value is not None else "")

            # 准备带软回车的错误占位符, 例如 '{{APP_NAME}}\v'
            broken_placeholder = key.replace("}}", f"}}{soft_break}")

            # 先替换可能存在的错误版本，再替换正常的版本
            paragraph.text = paragraph.text.replace(broken_placeholder, value_str)
            paragraph.text = paragraph.text.replace(key, value_str)

        # 防止因无法匹配的占位符导致的死循环
        if original_text == paragraph.text:
            if logging.getLogger().isEnabledFor(logging.WARNING):
                logging.warning(
                    "  > 警告: 在段落中发现无法匹配的占位符，已跳过。段落内容(repr): %s",
                    repr(paragraph.text[:90]),
                )
            break


def _parse_date(value: str | None) -> date | None:
    if not value:
        return None
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None


def _to_workday(value: date) -> date:
    while value.weekday() >= 5:
        value = value - timedelta(days=1)
    return value


def _to_dot(value: bool) -> str:
    return CHECKED_DOT if value else UNCHECKED_DOT


def _to_symbol(value: bool) -> str:
    return CHECKED_SYMBOL if value else UNCHECKED_SYMBOL


def _safe_str(value: object) -> str:
    if value is None:
        return ""
    return str(value)


def _get_first_holder(profile: dict) -> dict:
    holders = profile.get("copyright_holders", [])
    if holders:
        return holders[0]
    return {}


def build_copyright_replacements(company_profile: dict, data: dict) -> dict:
    replacements: dict[str, str] = {}
    def set_value(upper_key: str, lower_key: str, value: object) -> None:
        value_str = _safe_str(value)
        replacements[upper_key] = value_str
        replacements[lower_key] = value_str

    app_name = data.get("app__name", "")
    app_version = data.get("app__version", "")
    app_short_name = data.get("app__short_name", "")
    app_classification_code = data.get("app__classification_code", "")
    tech_source_lines = data.get("tech__source_lines", "")

    set_value("{{APP_NAME}}", "{{app__name}}", app_name)
    set_value("{{APP_VERSION}}", "{{app__version}}", app_version)
    set_value("{{APP_SHORT_NAME}}", "{{app__short_name}}", app_short_name)
    set_value("{{APP_CLASSIFICATION_CODE}}", "{{app__classification_code}}", app_classification_code)
    set_value("{{TECH_SOURCE_LINES}}", "{{tech__source_lines}}", tech_source_lines)

    completion_date = _parse_date(data.get("copyright__completion_date"))
    if completion_date is None:
        completion_date = _to_workday(date.today() - timedelta(days=14))
    set_value(
        "{{APP_COMPLETION_YEAR}}",
        "{{copyright__completion_year}}",
        completion_date.year if completion_date else "",
    )
    set_value(
        "{{APP_COMPLETION_MONTH}}",
        "{{copyright__completion_month}}",
        completion_date.month if completion_date else "",
    )
    set_value(
        "{{APP_COMPLETION_DAY}}",
        "{{copyright__completion_day}}",
        completion_date.day if completion_date else "",
    )

    published = bool(data.get("copyright__status_published", False))
    set_value("{{APP_STATUS_PUBLISHED}}", "{{copyright__status_published}}", _to_dot(published))
    set_value("{{APP_STATUS_UNPUBLISHED}}", "{{copyright__status_unpublished}}", _to_dot(not published))
    set_value(
        "{{APP_PUBLISH_DATE}}",
        "{{copyright__publish_date}}",
        data.get("copyright__publish_date", "") if published else "",
    )
    set_value(
        "{{APP_PUBLISH_LOCATION}}",
        "{{copyright__publish_location}}",
        data.get("copyright__publish_location", "") if published else "",
    )

    dev_method = data.get("copyright__development_method", "独立开发")
    set_value(
        "{{DEV_METHOD_INDEPENDENT}}",
        "{{copyright__dev_method_independent}}",
        _to_dot(dev_method == "独立开发"),
    )
    set_value(
        "{{DEV_METHOD_COOPERATIVE}}",
        "{{copyright__dev_method_cooperative}}",
        _to_dot(dev_method == "合作开发"),
    )
    set_value(
        "{{DEV_METHOD_COMMISSIONED}}",
        "{{copyright__dev_method_commissioned}}",
        _to_dot(dev_method == "委托开发"),
    )
    set_value(
        "{{DEV_METHOD_TASK_ASSIGNED}}",
        "{{copyright__dev_method_task_assigned}}",
        _to_dot(dev_method == "下达任务开发"),
    )

    acquire_method = data.get("rights__acquire_method", "原始取得")
    succession_types = {"受让", "承受", "继承"}
    is_succession = acquire_method in succession_types
    set_value(
        "{{RIGHTS_ACQUIRE_ORIGINAL}}",
        "{{rights__acquire_original}}",
        _to_dot(not is_succession),
    )
    set_value(
        "{{RIGHTS_ACQUIRE_SUCCESSION}}",
        "{{rights__acquire_succession}}",
        _to_dot(is_succession),
    )

    set_value(
        "{{RIGHTS_SUCCESSION_ASSIGNMENT}}",
        "{{rights__succession_assignment}}",
        _to_dot(acquire_method == "受让"),
    )
    set_value(
        "{{RIGHTS_SUCCESSION_ASSUMPTION}}",
        "{{rights__succession_assumption}}",
        _to_dot(acquire_method == "承受"),
    )
    set_value(
        "{{RIGHTS_SUCCESSION_INHERIT}}",
        "{{rights__succession_inherit}}",
        _to_dot(acquire_method == "继承"),
    )

    succession_details = data.get("rights__succession_details", {}) or {}
    set_value(
        "{{RIGHTS_SUCCESSION_IS_REGISTERED}}",
        "{{rights__succession_is_registered}}",
        _to_symbol(bool(succession_details.get("is_registered"))),
    )
    set_value(
        "{{RIGHTS_SUCCESSION_ORIGINAL_ID}}",
        "{{rights__succession_original_id}}",
        succession_details.get("original_id", ""),
    )
    set_value(
        "{{RIGHTS_SUCCESSION_IS_MODIFIED}}",
        "{{rights__succession_is_modified}}",
        _to_symbol(bool(succession_details.get("is_modified"))),
    )
    set_value(
        "{{RIGHTS_SUCCESSION_MODIFIED_CERT_ID}}",
        "{{rights__succession_modified_cert_id}}",
        succession_details.get("modified_cert_id", ""),
    )

    scope = data.get("rights__scope", "全部")
    set_value("{{RIGHTS_SCOPE_ALL}}", "{{rights__scope_all}}", _to_dot(scope == "全部"))
    set_value("{{RIGHTS_SCOPE_PARTIAL}}", "{{rights__scope_partial}}", _to_dot(scope == "部分"))

    partial_rights = data.get("rights__partial_rights", {}) or {}
    partial_keys = [
        "publish",
        "attribution",
        "modification",
        "copy",
        "distribution",
        "rental",
        "network",
        "translation",
        "other",
    ]
    for key in partial_keys:
        value = bool(partial_rights.get(key)) if scope == "部分" else False
        set_value(
            f"{{{{RIGHTS_PARTIAL_{key.upper()}}}}}",
            f"{{{{rights__partial_{key}}}}}",
            _to_symbol(value),
        )

    modification = data.get("modification_details") or {}
    is_modified = bool(modification)
    set_value("{{APP_TYPE_ORIGINAL}}", "{{copyright__app_type_original}}", _to_dot(not is_modified))
    set_value("{{APP_TYPE_MODIFIED}}", "{{copyright__app_type_modified}}", _to_dot(is_modified))
    set_value(
        "{{APP_MODIFIED_AUTH}}",
        "{{copyright__app_modified_auth}}",
        _to_symbol(bool(modification.get("authorized"))),
    )
    set_value(
        "{{APP_MODIFIED_REGISTERED}}",
        "{{copyright__app_modified_registered}}",
        _to_symbol(bool(modification.get("registered"))),
    )
    set_value(
        "{{APP_MODIFIED_ORIGINAL_ID}}",
        "{{copyright__app_modified_original_id}}",
        modification.get("original_id", ""),
    )
    set_value(
        "{{APP_MODIFIED_DESCRIPTION}}",
        "{{copyright__app_modified_description}}",
        modification.get("description", ""),
    )

    holders = company_profile.get("copyright_holders", [])
    holder_keys = ["name", "category", "id_type", "id_number", "nationality", "city", "found_date"]
    for idx in range(1, 5):
        if idx <= len(holders):
            holder = holders[idx - 1]
            for key in holder_keys:
                set_value(
                    f"{{{{AUTHOR_{idx}_{key.upper()}}}}}",
                    f"{{{{author_{idx}__{key}}}}}",
                    holder.get(key, ""),
                )
        else:
            for key in holder_keys:
                set_value(
                    f"{{{{AUTHOR_{idx}_{key.upper()}}}}}",
                    f"{{{{author_{idx}__{key}}}}}",
                    "",
                )

    applicant_source = company_profile.get("applicant_info") or _get_first_holder(company_profile)
    applicant_keys = [
        "name",
        "phone",
        "address",
        "zip_code",
        "contact_person",
        "mobile",
        "email",
        "fax",
    ]
    for key in applicant_keys:
        set_value(
            f"{{{{APPLICANT_{key.upper()}}}}}",
            f"{{{{applicant__{key}}}}}",
            applicant_source.get(key, ""),
        )

    applicant_type = str(data.get("applicant__type", "holder")).lower()
    is_agent = applicant_type in {"agent", "代理人", "代理"}
    set_value("{{APPLICANT_TYPE_HOLDER}}", "{{applicant__type_holder}}", _to_dot(not is_agent))
    set_value("{{APPLICANT_TYPE_AGENT}}", "{{applicant__type_agent}}", _to_dot(is_agent))
    set_value("{{DELEGATION_STATEMENT}}", "{{delegation_statement}}", "")
    for key in applicant_keys:
        set_value(
            f"{{{{AGENT_{key.upper()}}}}}",
            f"{{{{agent__{key}}}}}",
            "",
        )

    signature_date = _parse_date(data.get("signature__date")) or _to_workday(date.today())
    set_value("{{SIGNATURE_YEAR}}", "{{signature__year}}", signature_date.year)
    set_value("{{SIGNATURE_MONTH}}", "{{signature__month}}", signature_date.month)
    set_value("{{SIGNATURE_DAY}}", "{{signature__day}}", signature_date.day)

    tech_map = {
        "TECH_HARDWARE_DEV": data.get("tech__hardware_dev", ""),
        "TECH_HARDWARE_RUN": data.get("tech__hardware_run", ""),
        "TECH_OS_DEV": data.get("tech__os_dev", ""),
        "TECH_OS_RUN": data.get("tech__os_run", ""),
        "TECH_DEV_TOOLS": data.get("tech__dev_tools", ""),
        "TECH_RUN_SUPPORT": data.get("tech__run_support", ""),
        "TECH_LANGUAGE": data.get("tech__language", ""),
        "TECH_DEV_PURPOSE": data.get("tech__dev_purpose", ""),
        "TECH_MAIN_FUNCTIONS": data.get("tech__main_functions", ""),
        "TECH_FEATURES": data.get("tech__features", ""),
    }
    tech_lower_map = {
        "TECH_HARDWARE_DEV": "{{tech__hardware_dev}}",
        "TECH_HARDWARE_RUN": "{{tech__hardware_run}}",
        "TECH_OS_DEV": "{{tech__os_dev}}",
        "TECH_OS_RUN": "{{tech__os_run}}",
        "TECH_DEV_TOOLS": "{{tech__dev_tools}}",
        "TECH_RUN_SUPPORT": "{{tech__run_support}}",
        "TECH_LANGUAGE": "{{tech__language}}",
        "TECH_DEV_PURPOSE": "{{tech__dev_purpose}}",
        "TECH_MAIN_FUNCTIONS": "{{tech__main_functions}}",
        "TECH_FEATURES": "{{tech__features}}",
    }
    for key, value in tech_map.items():
        lower_key = tech_lower_map.get(key, f"{{{{{key.lower()}}}}}")
        set_value(f"{{{{{key}}}}}", lower_key, value)

    return replacements


def generate_document(company_profile: dict, data: dict, template_path: Path, output_path: Path) -> None:
    """
    根据公司信息与统一字段数据生成软件著作权登记申请表。
    """
    replacements = build_copyright_replacements(company_profile, data)

    try:
        doc = Document(template_path)
    except Exception as exc:
        logging.error("❌ 错误: 无法打开Word模板文件 '%s': %s", template_path, exc)
        return

    docx_replace_text(doc, replacements)

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"生成文件: {output_path}")
    except Exception as exc:
        logging.error("❌ 错误: 保存文件失败 '%s': %s", output_path, exc)


if __name__ == "__main__":
    print("This module is intended to be used via the CLI.")
