from __future__ import annotations

from typing import Dict


def _replace_text(text: str, placeholder_map: Dict[str, str]) -> str:
    for key in sorted(placeholder_map.keys(), key=len, reverse=True):
        value = placeholder_map[key]
        if key in text:
            text = text.replace(key, value)
    return text


def replace_in_paragraph(paragraph, placeholder_map: Dict[str, str]) -> None:
    if not paragraph.runs:
        if paragraph.text:
            paragraph.text = _replace_text(paragraph.text, placeholder_map)
        return
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = _replace_text(full_text, placeholder_map)
    if replaced == full_text:
        return
    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_in_table(table, placeholder_map: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, placeholder_map)


def replace_in_header_footer(doc, placeholder_map: Dict[str, str]) -> None:
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, placeholder_map)
        for table in header.tables:
            replace_in_table(table, placeholder_map)
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, placeholder_map)
        for table in footer.tables:
            replace_in_table(table, placeholder_map)


def _replace_in_part_xml(part, placeholder_map: Dict[str, str]) -> None:
    element = part._element
    textbox_ps = element.xpath(".//w:txbxContent//w:p")
    for p in list(textbox_ps):
        texts = p.xpath(".//w:t")
        if not texts:
            continue
        full_text = "".join(t.text or "" for t in texts)
        if "{{" not in full_text or "}}" not in full_text:
            continue
        replaced = _replace_text(full_text, placeholder_map)
        if replaced == full_text:
            continue
        texts[0].text = replaced
        for t in texts[1:]:
            t.text = ""


def fill_docx(template_path: str, placeholder_map: Dict[str, str], out_path: str) -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise ImportError("Missing dependency: python-docx") from exc

    doc = Document(template_path)
    _replace_in_part_xml(doc.part, placeholder_map)
    for section in doc.sections:
        _replace_in_part_xml(section.header.part, placeholder_map)
        _replace_in_part_xml(section.footer.part, placeholder_map)
    doc.save(out_path)
