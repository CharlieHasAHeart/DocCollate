from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


def _replace(doc: Document, replacements: dict[str, str]) -> None:
    def replace_in_paragraph(p):
        while "{{" in p.text:
            old = p.text
            for k, v in replacements.items():
                p.text = p.text.replace(k, v)
                p.text = p.text.replace(k.replace("}}", "}}\\v"), v)
            if old == p.text:
                break

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)


def generate_document(template_path: Path, output_path: Path, data: dict[str, str]) -> None:
    try:
        doc = Document(template_path)
    except Exception as exc:
        logger.error("Unable to open template '%s': %s", template_path, exc)
        raise

    replacements = {
        "{{app__name}}": data.get("app__name", ""),
        "{{app__short_name}}": data.get("app__short_name", ""),
        "{{app__version}}": data.get("app__version", ""),
        "{{env__dev_lang}}": data.get("env__dev_lang", ""),
        "{{env__dev_platform}}": data.get("env__dev_platform", ""),
        "{{env__run_platform}}": data.get("env__run_platform", ""),
        "{{product__app_domain}}": data.get("product__app_domain", ""),
        "{{holder__name}}": data.get("holder__name", ""),
        "{{holder__address}}": data.get("holder__address", ""),
        "{{holder__zip_code}}": data.get("holder__zip_code", ""),
        "{{holder__contact_name}}": data.get("holder__contact_name", ""),
        "{{holder__contact_mobile}}": data.get("holder__contact_mobile", ""),
        "{{holder__contact_email}}": data.get("holder__contact_email", ""),
        "{{holder__contact_landline}}": data.get("holder__contact_landline", ""),
        "{{holder__tech_contact_name}}": data.get("holder__tech_contact_name", ""),
        "{{holder__tech_contact_mobile}}": data.get("holder__tech_contact_mobile", ""),
    }
    _replace(doc, replacements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
