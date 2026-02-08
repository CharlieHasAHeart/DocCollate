from __future__ import annotations

import logging
from pathlib import Path
from types import SimpleNamespace

from docx import Document

logger = logging.getLogger(__name__)


def _to_rows(module_list: list[dict]) -> list[SimpleNamespace]:
    rows: list[SimpleNamespace] = []
    for module in module_list:
        if not isinstance(module, dict):
            continue
        module_name = str(module.get("name", "")).strip()
        if not module_name:
            continue

        items_raw = module.get("items", [])
        items: list[SimpleNamespace] = []
        if isinstance(items_raw, list):
            for item in items_raw:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("name", "")).strip()
                desc = str(item.get("desc", "")).strip()
                if not (name or desc):
                    continue
                items.append((name, desc))

        for level2, desc in items:
            rows.append(
                SimpleNamespace(
                    level1=module_name,
                    level2=level2,
                    desc=desc,
                )
            )
    return rows


def _set_cell_text(cell, text: str) -> None:
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    para = cell.paragraphs[0]
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    para.text = str(text or "")


def _render_rows_table(template_path: Path, output_path: Path, rows: list[SimpleNamespace]) -> None:
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template has no table")
    table = doc.tables[0]
    if len(table.columns) < 3:
        raise ValueError("Template first table must have at least 3 columns")

    # Keep only header and one data prototype row.
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)
    if len(table.rows) < 2:
        table.add_row()

    if rows:
        first = rows[0]
        _set_cell_text(table.rows[1].cells[0], first.level1)
        _set_cell_text(table.rows[1].cells[1], first.level2)
        _set_cell_text(table.rows[1].cells[2], first.desc)

        for r in rows[1:]:
            row = table.add_row()
            _set_cell_text(row.cells[0], r.level1)
            _set_cell_text(row.cells[1], r.level2)
            _set_cell_text(row.cells[2], r.desc)
    else:
        _set_cell_text(table.rows[1].cells[0], "")
        _set_cell_text(table.rows[1].cells[1], "")
        _set_cell_text(table.rows[1].cells[2], "")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_document(template_path: Path, output_path: Path, module_list: list[dict]) -> None:
    rows = _to_rows(module_list)
    logger.info("Renderer mode=manual_rows file=%s rows=%s", __file__, len(rows))
    _render_rows_table(template_path, output_path, rows)
