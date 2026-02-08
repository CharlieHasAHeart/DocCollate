from __future__ import annotations

import re
from typing import Any


_BULLET_RE = re.compile(r"^\s*(?:- |\u2022 )(.+?)\s*$")
_NUMBER_RE = re.compile(r"^\s*\d+(?:[.)]|、)\s+(.+?)\s*$")


def _coerce_map(placeholder_map: dict[str, Any]) -> dict[str, str]:
    out: dict[str, str] = {}
    for k, v in placeholder_map.items():
        if v is None:
            out[str(k)] = ""
        elif isinstance(v, str):
            out[str(k)] = v
        else:
            out[str(k)] = str(v)
    return out


def _replace_text(text: str, placeholder_map: dict[str, str]) -> str:
    """
    Replace placeholders in a stable order (longest keys first).
    """
    for key in sorted(placeholder_map.keys(), key=len, reverse=True):
        if key in text:
            text = text.replace(key, placeholder_map[key])
    return text


def _copy_paragraph_format(src, dst) -> None:
    try:
        src_fmt = src.paragraph_format
        dst_fmt = dst.paragraph_format
    except Exception:
        return
    attrs = [
        "alignment",
        "first_line_indent",
        "left_indent",
        "right_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacing_rule",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    ]
    for attr in attrs:
        try:
            value = getattr(src_fmt, attr)
            setattr(dst_fmt, attr, value)
        except Exception:
            continue


def _insert_paragraph_after(
    paragraph,
    text: str = "",
    style: str | None = None,
    template_paragraph=None,
):
    """
    Insert a new paragraph directly after `paragraph`.
    Works in normal document body and table cells.
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if template_paragraph is not None:
        if not style:
            try:
                new_para.style = template_paragraph.style
            except Exception:
                pass
        if not style or str(style).lower() not in {"list bullet", "list number"}:
            _copy_paragraph_format(template_paragraph, new_para)
    if text:
        new_para.add_run(text)
    return new_para


def _clear_paragraph(paragraph) -> None:
    """
    Remove all runs in a paragraph (leave the paragraph itself).
    """
    for run in list(paragraph.runs):
        try:
            run.text = ""
        except Exception:
            pass


def _render_rich_text_into_paragraph(paragraph, value: str) -> None:
    """
    Render multiline value into docx as real paragraphs / lists.

    Rules:
    - Blank-line (`\n\n`) separates paragraphs/blocks
    - Lines starting with '- ' or '• ' become bullet list items (List Bullet)
    - Lines starting with '1. ' / '2) ' / '3、 ' become numbered list items (List Number)
    - Other lines become normal paragraphs (single '\n' is kept as line break within paragraph)
    """
    blocks = value.replace("\r\n", "\n").replace("\r", "\n").split("\n\n")
    blocks = [b for b in (blk.strip("\n") for blk in blocks) if b.strip()]
    if not blocks:
        _clear_paragraph(paragraph)
        return

    # We will re-use the original paragraph as the "anchor".
    anchor = paragraph
    _clear_paragraph(anchor)

    first_written = False
    cursor = anchor

    def write_normal_para(target_para, text: str) -> None:
        # Keep single '\n' as line breaks inside the paragraph.
        # python-docx will translate '\n' into a line break in a run.
        target_para.add_run(text)

    for block in blocks:
        lines = [ln.rstrip() for ln in block.split("\n")]
        lines = [ln for ln in lines if ln.strip() != ""]
        if not lines:
            continue

        # Decide if this block is a pure list (all lines match list markers)
        bullet_items: list[str] = []
        number_items: list[str] = []
        mixed = False
        for line in lines:
            bullet_match = _BULLET_RE.match(line)
            number_match = _NUMBER_RE.match(line)
            if bullet_match:
                bullet_items.append(bullet_match.group(1))
            elif number_match:
                number_items.append(number_match.group(1))
            else:
                mixed = True
                break

        if not mixed and bullet_items and not number_items:
            # Bullet list block
            for item in bullet_items:
                if not first_written:
                    try:
                        cursor.style = "List Bullet"
                    except Exception:
                        pass
                    write_normal_para(cursor, item)
                    first_written = True
                else:
                    cursor = _insert_paragraph_after(cursor, item, style="List Bullet")
            continue

        if not mixed and number_items and not bullet_items:
            # Numbered list block
            for item in number_items:
                if not first_written:
                    try:
                        cursor.style = "List Number"
                    except Exception:
                        pass
                    write_normal_para(cursor, item)
                    first_written = True
                else:
                    cursor = _insert_paragraph_after(cursor, item, style="List Number")
            continue

        # Mixed/normal text block: join lines back with '\n' (line breaks)
        joined = "\n".join(lines).strip()
        if not first_written:
            write_normal_para(cursor, joined)
            first_written = True
        else:
            cursor = _insert_paragraph_after(cursor, joined, style=None, template_paragraph=anchor)


def replace_in_paragraph(paragraph, placeholder_map: dict[str, str]) -> None:
    if not paragraph.runs:
        if paragraph.text:
            paragraph.text = _replace_text(paragraph.text, placeholder_map)
        return
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = _replace_text(full_text, placeholder_map)
    if replaced == full_text:
        return
    # If the replacement introduces structure, render it as real paragraphs/lists.
    if "\n" in replaced:
        _render_rich_text_into_paragraph(paragraph, replaced)
        return
    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_in_table(table, placeholder_map: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, placeholder_map)


def replace_in_header_footer(doc, placeholder_map: dict[str, str]) -> None:
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, placeholder_map)
        for table in header.tables:
            replace_in_table(table, placeholder_map)
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, placeholder_map)
        for table in footer.tables:
            replace_in_table(table, placeholder_map)


def _replace_in_part_xml(part, placeholder_map: dict[str, str]) -> None:
    """
    Fallback replacer for textboxes/shapes where python-docx doesn't expose
    a friendly paragraph API. This is plain text replacement (no list/paragraph insert).
    """
    element = part._element
    textbox_ps = element.xpath(".//w:txbxContent//w:p")
    for p in list(textbox_ps):
        texts = p.xpath(".//w:t")
        if not texts:
            continue
        full_text = "".join(t.text or "" for t in texts)
        if "{{" not in full_text or "}}" not in full_text:
            continue
        replaced = _replace_text(full_text, placeholder_map)
        if replaced == full_text:
            continue
        texts[0].text = replaced
        for t in texts[1:]:
            t.text = ""


def fill_docx(template_path: str, placeholder_map: dict[str, Any], out_path: str) -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise ImportError("Missing dependency: python-docx") from exc

    placeholder_map_s = _coerce_map(placeholder_map)
    doc = Document(template_path)

    # 1) Normal document body (paragraphs + tables)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, placeholder_map_s)
    for table in doc.tables:
        replace_in_table(table, placeholder_map_s)

    # 2) Header/Footer (paragraphs + tables)
    replace_in_header_footer(doc, placeholder_map_s)

    # 3) Textboxes / shapes via raw XML fallback
    _replace_in_part_xml(doc.part, placeholder_map_s)
    for section in doc.sections:
        _replace_in_part_xml(section.header.part, placeholder_map_s)
        _replace_in_part_xml(section.footer.part, placeholder_map_s)

    doc.save(out_path)
