from __future__ import annotations

import json
import logging
import os
import re
from pathlib import Path
from typing import Any

from proposal_app.config import AppConfig, load_config, load_dotenv
from proposal_app.llm.api import translate_to_english
from proposal_app.llm.client import LLMRuntime, init_llm
from proposal_app.proposal.cluster_defs import PLACEHOLDER_FIELDS, TABLE_MIN_SPECS
from proposal_app.proposal.inputs import prompt_cover
from proposal_app.proposal.mapping import build_placeholder_map
from proposal_app.proposal.spec_loader import load_spec_text
from proposal_app.proposal.utils import ensure_dir
from proposal_app.render.api import render_docx_from_output

from .graph import build_graph


_ROOT = Path(__file__).resolve().parents[2]
_TEMPLATE_SECTIONS_JSON = _ROOT / "debug" / "template_sections.json"

# Ensure local .env is loaded early so PROPOSAL_CONFIG/DEBUG flags can be honored.
load_dotenv(_ROOT / ".env")

logger = logging.getLogger(__name__)


def _resolve_config_path(config_path: str | None) -> str:
    if config_path:
        return config_path
    env_config = os.getenv("PROPOSAL_CONFIG", "").strip()
    if env_config:
        return env_config
    return str(_ROOT / "pyproject.toml")


def _init_runtimes(app_config: AppConfig, args: Any) -> tuple[LLMRuntime, LLMRuntime, LLMRuntime]:
    runtime = init_llm(
        app_config.llm,
        api_key=getattr(args, "api_key", None),
        base_url=getattr(args, "base_url", None),
        model=getattr(args, "model", None),
    )
    ledger_model = getattr(args, "skeleton_model", None) or app_config.llm.skeleton_model or runtime.model
    final_model = getattr(args, "final_model", None) or app_config.llm.final_model or runtime.model
    ledger_runtime = LLMRuntime(client=runtime.client, model=ledger_model, api_key=runtime.api_key, base_url=runtime.base_url)
    final_runtime = LLMRuntime(client=runtime.client, model=final_model, api_key=runtime.api_key, base_url=runtime.base_url)
    return runtime, ledger_runtime, final_runtime


_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")


def _heading_level(style_id_or_name: str) -> int | None:
    s = (style_id_or_name or "").strip().lower().replace("_", " ")
    s = re.sub(r"\s+", " ", s)
    if not s:
        return None
    if "heading" in s:
        if "1" in s:
            return 1
        if "2" in s:
            return 2
    if "标题" in style_id_or_name:
        if "1" in style_id_or_name:
            return 1
        if "2" in style_id_or_name:
            return 2
    if s.endswith("1"):
        return 1
    if s.endswith("2"):
        return 2
    return None


def _extract_required_placeholders(template_path: Path) -> list[str]:
    if not template_path.exists():
        return []

    allowed = set(PLACEHOLDER_FIELDS)

    def _norm(tag: str) -> str:
        inner = tag.strip()
        inner = inner.lstrip("{").rstrip("}")
        inner = inner.strip()
        return "{{ " + inner + " }}"

    def _iter_paragraph_texts(doc: Any) -> list[str]:
        texts: list[str] = []
        for p in getattr(doc, "paragraphs", []):
            runs = getattr(p, "runs", None)
            if runs:
                texts.append("".join(r.text or "" for r in runs))
            else:
                texts.append(getattr(p, "text", "") or "")
        for t in getattr(doc, "tables", []):
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        runs = getattr(p, "runs", None)
                        if runs:
                            texts.append("".join(r.text or "" for r in runs))
                        else:
                            texts.append(getattr(p, "text", "") or "")
        return texts

    seq: list[str] = []
    try:
        from docx import Document

        doc = Document(str(template_path))
        for text in _iter_paragraph_texts(doc):
            for m in _PLACEHOLDER_RE.finditer(text):
                tag = _norm(m.group(0))
                if tag in allowed:
                    seq.append(tag)
        for section in doc.sections:
            for part in [section.header, section.footer]:
                for text in _iter_paragraph_texts(part):
                    for m in _PLACEHOLDER_RE.finditer(text):
                        tag = _norm(m.group(0))
                        if tag in allowed:
                            seq.append(tag)
    except Exception:
        import html
        import zipfile

        def _iter_docx_xml_texts(path: Path) -> list[str]:
            texts: list[str] = []
            with zipfile.ZipFile(path, "r") as z:
                names = [n for n in z.namelist() if n.startswith("word/")]
                names.sort()
                for name in names:
                    if not (name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer")):
                        continue
                    xml = z.read(name).decode("utf-8", errors="ignore")
                    for m in re.finditer(r"<w:t[^>]*>(.*?)</w:t>", xml):
                        texts.append(html.unescape(m.group(1)))
            return texts

        whole = "".join(_iter_docx_xml_texts(template_path))
        for m in _PLACEHOLDER_RE.finditer(whole):
            tag = _norm(m.group(0))
            if tag in allowed:
                seq.append(tag)

    start = "{{ purpose }}"
    end = PLACEHOLDER_FIELDS[-1] if PLACEHOLDER_FIELDS else "{{ purpose }}"
    try:
        s_idx = seq.index(start)
        e_idx = seq.index(end)
    except ValueError:
        return []
    if s_idx > e_idx:
        return []

    out: list[str] = []
    seen: set[str] = set()
    for tag in seq[s_idx : e_idx + 1]:
        if tag in seen:
            continue
        seen.add(tag)
        out.append(tag)
    return out


def _extract_placeholder_sections_by_heading(template_path: Path) -> list[list[str]]:
    if not template_path.exists():
        return []

    allowed = set(PLACEHOLDER_FIELDS)

    def _norm(tag: str) -> str:
        inner = tag.strip()
        inner = inner.lstrip("{").rstrip("}")
        inner = inner.strip()
        return "{{ " + inner + " }}"

    start = "{{ purpose }}"
    end = PLACEHOLDER_FIELDS[-1] if PLACEHOLDER_FIELDS else "{{ purpose }}"

    sections: list[list[str]] = []
    cur: list[str] = []
    seen_in_cur: set[str] = set()
    started = False
    ended = False

    def _flush() -> None:
        nonlocal cur, seen_in_cur
        if cur:
            sections.append(cur)
        cur = []
        seen_in_cur = set()

    def _push_tag(tag: str) -> None:
        nonlocal started, ended
        if not started:
            if tag != start:
                return
            started = True
        if ended:
            return
        if tag not in seen_in_cur:
            cur.append(tag)
            seen_in_cur.add(tag)
        if tag == end:
            ended = True

    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(str(template_path))

        def _iter_blocks(parent: Any) -> list[tuple[str, Any]]:
            out: list[tuple[str, Any]] = []
            body = getattr(parent, "element", None)
            if body is None:
                return out
            body_child = getattr(body, "body", None)
            container = body_child if body_child is not None else body
            for child in container.iterchildren():
                if isinstance(child, CT_P):
                    out.append(("p", Paragraph(child, parent)))
                elif isinstance(child, CT_Tbl):
                    out.append(("tbl", Table(child, parent)))
            return out

        for kind, block in _iter_blocks(doc):
            if ended:
                break
            if kind == "p":
                style_name = ""
                try:
                    style_name = str(getattr(block.style, "name", "") or getattr(block.style, "style_id", "") or "")
                except Exception:
                    style_name = ""
                lvl = _heading_level(style_name)
                if started and not ended and lvl in (1, 2):
                    _flush()
                text = "".join((r.text or "") for r in getattr(block, "runs", [])).strip() or (getattr(block, "text", "") or "")
                for m in _PLACEHOLDER_RE.finditer(text):
                    tag = _norm(m.group(0))
                    if tag in allowed:
                        _push_tag(tag)
                        if ended:
                            break
            else:
                for row in block.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            text = "".join((r.text or "") for r in getattr(p, "runs", [])).strip() or (getattr(p, "text", "") or "")
                            for m in _PLACEHOLDER_RE.finditer(text):
                                tag = _norm(m.group(0))
                                if tag in allowed:
                                    _push_tag(tag)
                                    if ended:
                                        break
                            if ended:
                                break
                        if ended:
                            break
                    if ended:
                        break
    except Exception:
        import html
        import zipfile

        with zipfile.ZipFile(template_path, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        for pm in re.finditer(r"<w:p\b[\s\S]*?</w:p>", xml):
            if ended:
                break
            p_xml = pm.group(0)
            style_id = ""
            sm = re.search(r"<w:pStyle[^>]*w:val=\"([^\"]+)\"", p_xml)
            if sm:
                style_id = sm.group(1)
            lvl = _heading_level(style_id)
            if started and not ended and lvl in (1, 2):
                _flush()
            text = "".join(html.unescape(x) for x in re.findall(r"<w:t[^>]*>(.*?)</w:t>", p_xml))
            for m in _PLACEHOLDER_RE.finditer(text):
                tag = _norm(m.group(0))
                if tag in allowed:
                    _push_tag(tag)
                    if ended:
                        break

    if started:
        _flush()
    return sections


def _extract_sections_from_template_json(path: Path, chunk_size: int) -> tuple[list[str], list[list[str]], list[list[str]]]:
    if not path.exists():
        return [], [], []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return [], [], []
    chapters = data.get("chapters", [])
    if not isinstance(chapters, list):
        return [], [], []

    allowed = set(PLACEHOLDER_FIELDS)

    def _norm(tag: str) -> str:
        inner = tag.strip()
        inner = inner.lstrip("{").rstrip("}")
        inner = inner.strip()
        return "{{ " + inner + " }}"

    section_placeholders: list[list[str]] = []
    table_order: list[str] = []
    seen_tables: set[str] = set()

    def _collect_items(items: Any) -> list[str]:
        placeholders: list[str] = []
        if not isinstance(items, list):
            return placeholders
        for item in items:
            if isinstance(item, str):
                tag = _norm(item)
                if tag in allowed:
                    placeholders.append(tag)
            elif isinstance(item, dict) and "table" in item:
                table_name = str(item.get("table", "")).strip()
                if table_name and table_name not in seen_tables:
                    seen_tables.add(table_name)
                    table_order.append(table_name)
        return placeholders

    for chapter in chapters:
        if not isinstance(chapter, dict):
            continue
        sections = chapter.get("sections", [])
        if isinstance(sections, list) and sections:
            for section in sections:
                if not isinstance(section, dict):
                    continue
                sec_ph = _collect_items(section.get("placeholders", []))
                if sec_ph:
                    section_placeholders.append(sec_ph)
            continue
        # Fallback: chapter-level placeholders if sections are missing.
        chap_ph = _collect_items(chapter.get("placeholders", []))
        if chap_ph:
            section_placeholders.append(chap_ph)

    required_placeholders: list[str] = []
    seen_ph: set[str] = set()
    for sec in section_placeholders:
        for tag in sec:
            if tag in seen_ph:
                continue
            seen_ph.add(tag)
            required_placeholders.append(tag)

    required_chunks: list[list[str]] = []
    for sec in section_placeholders:
        if chunk_size > 0 and len(sec) > chunk_size:
            required_chunks.extend(_chunk_placeholders(sec, chunk_size=chunk_size))
        else:
            required_chunks.append(sec)

    table_chunks = [[name] for name in table_order]
    return required_placeholders, required_chunks, table_chunks


def _chunk_placeholders(placeholders: list[str], chunk_size: int) -> list[list[str]]:
    if chunk_size <= 0:
        chunk_size = 10
    chunks: list[list[str]] = []
    cur: list[str] = []
    for p in placeholders:
        if not isinstance(p, str) or not p.strip():
            continue
        cur.append(p)
        if len(cur) >= chunk_size:
            chunks.append(cur)
            cur = []
    if cur:
        chunks.append(cur)
    return chunks


def _normalize_manual_inputs(manual_inputs: dict[str, Any]) -> dict[str, Any]:
    return manual_inputs


def _apply_schedule_inputs(manual_inputs: dict[str, Any], args: Any) -> dict[str, Any]:
    start_date = getattr(args, "start_date", None) or str(manual_inputs.get("start_date", "")).strip()
    end_date = getattr(args, "end_date", None) or str(manual_inputs.get("end_date", "")).strip()
    if start_date:
        manual_inputs["start_date"] = start_date
    if end_date:
        manual_inputs["end_date"] = end_date
    return manual_inputs


def _load_manual_inputs(args: Any, app_config: AppConfig, runtime: LLMRuntime) -> dict[str, Any]:
    manual_inline = getattr(args, "manual_inputs", None)
    if isinstance(manual_inline, dict):
        manual_inputs = manual_inline
        if "cover" not in manual_inputs:
            company_name = getattr(args, "company_name", "").strip()
            project_name = getattr(args, "project_name", "") or ""
            if project_name:
                try:
                    english_name = translate_to_english(project_name, runtime)
                except Exception:
                    english_name = ""
                if not english_name or re.search(r"[\u4e00-\u9fff]", english_name):
                    english_name = project_name
            else:
                english_name = "Project"
            manual_inputs["cover"] = {
                "company_name": company_name,
                "project_name": project_name,
                "project_id": "",
                "document_title": f"{english_name} Project Proposal",
                "document_version": "V1.0",
                "drafted_by": "",
                "draft_date": "",
                "approved_by": "",
                "approval_date": "",
            }
        positioning = getattr(args, "positioning", "").strip()
        if positioning and "positioning" not in manual_inputs:
            manual_inputs["positioning"] = positioning
        manual_inputs = _normalize_manual_inputs(manual_inputs)
        manual_inputs = _apply_schedule_inputs(manual_inputs, args)
        return manual_inputs

    company_name = getattr(args, "company_name", "").strip()
    manual_inputs = {
        "cover": prompt_cover(
            runtime,
            company_name_default=company_name,
            project_name_default=getattr(args, "project_name", ""),
        ),
    }
    # Prompt schedule dates in interactive mode if not provided via args.
    manual_inputs = _normalize_manual_inputs(manual_inputs)
    positioning = getattr(args, "positioning", "").strip()
    if positioning:
        manual_inputs["positioning"] = positioning
    manual_inputs = _apply_schedule_inputs(manual_inputs, args)
    return manual_inputs


def _write_debug(
    debug_dir: str,
    llm_output: dict[str, Any],
    placeholder_map: dict[str, Any],
    ledger: dict[str, Any] | None,
    metrics: dict[str, Any] | None = None,
) -> None:
    ensure_dir(debug_dir)
    with open(os.path.join(debug_dir, "ledger_output.json"), "w", encoding="utf-8") as f:
        json.dump(ledger or {}, f, ensure_ascii=False, indent=2)
    with open(os.path.join(debug_dir, "llm_output.json"), "w", encoding="utf-8") as f:
        json.dump(llm_output, f, ensure_ascii=False, indent=2)
    with open(os.path.join(debug_dir, "placeholder_map.json"), "w", encoding="utf-8") as f:
        json.dump(placeholder_map, f, ensure_ascii=False, indent=2)
    if metrics:
        with open(os.path.join(debug_dir, "metrics.json"), "w", encoding="utf-8") as f:
            json.dump(metrics, f, ensure_ascii=False, indent=2)


def _write_metrics_log(metrics: dict[str, Any], out_dir: str) -> None:
    ensure_dir(out_dir)
    log_path = os.path.join(out_dir, "runs.jsonl")
    import time
    entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        **metrics
    }
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


def _print_metrics_summary(metrics: dict[str, Any]) -> None:
    logger.info("%s", "\n" + "=" * 20 + " Metrics Summary " + "=" * 20)
    if "llm_calls" in metrics:
        logger.info("LLM Calls:        %s", metrics.get("llm_calls"))
    logger.info("Gate First Pass:  %s", metrics.get("gate_first_pass"))
    logger.info("Gate Rounds:      %s (Repairs: %s)", metrics.get("gate_rounds"), metrics.get("gate_repair_count", 0))
    
    logger.info("Post First Pass:  %s", metrics.get("post_first_pass"))
    logger.info("Rewrite Rounds:   %s (Repairs: %s)", metrics.get("rewrite_rounds"), metrics.get("rewrite_repair_count", 0))
    if "patch_rounds" in metrics:
        logger.info("Patch Rounds:     %s", metrics.get("patch_rounds"))
        
    first_post = metrics.get('issues_by_rule_first_post', {})
    logger.info(
        "Post Lint Hits (First Pass): R4=%s, R5=%s, R6=%s, R7=%s",
        first_post.get("R4", 0),
        first_post.get("R5", 0),
        first_post.get("R6", 0),
        first_post.get("R7", 0),
    )

    final_issues = metrics.get('final_issues_by_rule', {})
    logger.info(
        "Final Lint Counts (Residual): R4=%s, R5=%s, R6=%s, R7=%s",
        final_issues.get("R4", 0),
        final_issues.get("R5", 0),
        final_issues.get("R6", 0),
        final_issues.get("R7", 0),
    )
    
    has_l0_fail = any(final_issues.get(r, 0) > 0 for r in ["R4", "R5", "R6", "R7"])
    logger.info("Final L0 Status:  %s", "FAIL" if has_l0_fail else "PASS")
        
    doc_chars = metrics.get('doc_char_count', 0)
    doc_k = doc_chars / 1000.0 if doc_chars > 0 else 1.0
    subj_hits = round(metrics.get('subjective_density_per_k', 0) * doc_k)
    num_hits = round(metrics.get('new_number_risk_per_k', 0) * doc_k)
    
    logger.info(
        "Subjective Density: %.2f /k (%s hits / %s chars)",
        metrics.get("subjective_density_per_k", 0),
        subj_hits,
        doc_chars,
    )
    logger.info(
        "New Number Risk:    %.2f /k (%s hits / %s chars)",
        metrics.get("new_number_risk_per_k", 0),
        num_hits,
        doc_chars,
    )
    logger.info("%s", "=" * 57 + "\n")


def _render_docx(
    app_config: AppConfig,
    manual_inputs: dict[str, Any],
    llm_output: dict[str, Any],
    placeholder_map: dict[str, Any],
    out_dir: str,
) -> str:
    template_path = app_config.templates.proposal
    if not template_path:
        raise ValueError("templates.proposal missing in config")
    os.makedirs(out_dir, exist_ok=True)
    project_name = placeholder_map.get("{{ project_name }}", "").strip()
    if not project_name:
        cover = manual_inputs.get("cover", {}) if isinstance(manual_inputs, dict) else {}
        project_name = str(cover.get("project_name", "")).strip()
    if not project_name:
        raise ValueError("project_name missing for output file name")
    out_path = os.path.join(out_dir, f"{project_name}立项建议书.docx")
    render_docx_from_output(
        template_path=str(template_path),
        out_path=out_path,
        llm_output=llm_output,
        manual_inputs=manual_inputs,
    )
    return out_path


def run_pipeline(args: Any) -> dict[str, Any]:
    debug_env = os.getenv("PROPOSAL_DEBUG", "").strip().lower()
    if debug_env in {"0", "false", "no", "n"}:
        debug_flag = False
    elif debug_env in {"1", "true", "yes", "y"}:
        debug_flag = True
    else:
        debug_arg = getattr(args, "debug", None)
        debug_flag = True if debug_arg is None else bool(debug_arg)
    debug_dir = getattr(args, "debug_dir", "") or os.getenv("PROPOSAL_DEBUG_DIR", "").strip() or "debug"

    llm_output: dict[str, Any] = {}
    placeholder_map: dict[str, Any] = {}
    ledger: dict[str, Any] | None = None
    metrics: dict[str, Any] | None = None
    manual_inputs: dict[str, Any] = {}
    stage = "init"
    error_info: dict[str, Any] | None = None

    try:
        stage = "load_config"
        config_path = _resolve_config_path(getattr(args, "config", None))
        app_config = load_config(config_path)

        stage = "init_runtimes"
        runtime, ledger_runtime, final_runtime = _init_runtimes(app_config, args)

        stage = "load_spec"
        spec_text = load_spec_text(args.spec)
        manual_inputs = _load_manual_inputs(args, app_config, runtime)

        stage = "prepare_placeholders"
        required_placeholders: list[str] = []
        chunk_size = 10
        placeholder_chunks: list[list[str]] = []
        table_chunks: list[list[str]] = []
        required_tables: list[str] = []
        if app_config.templates.proposal:
            if _TEMPLATE_SECTIONS_JSON.exists():
                req_ph, req_chunks, tbl_chunks = _extract_sections_from_template_json(_TEMPLATE_SECTIONS_JSON, chunk_size)
                if req_chunks:
                    logger.info("[Info] Using section order from debug/template_sections.json")
                    required_placeholders = req_ph
                    placeholder_chunks = req_chunks
                    table_chunks = tbl_chunks
                    # Append any placeholders missing from the JSON order to keep completeness checks honest.
                    extra = _extract_required_placeholders(app_config.templates.proposal)
                    if extra:
                        seen_extra = set(required_placeholders)
                        extra = [p for p in extra if p not in seen_extra]
                        if extra:
                            required_placeholders.extend(extra)
            if not required_placeholders:
                sections = _extract_placeholder_sections_by_heading(app_config.templates.proposal)
                if sections:
                    seen: set[str] = set()
                    for sec in sections:
                        if chunk_size > 0 and len(sec) > chunk_size:
                            placeholder_chunks.extend(_chunk_placeholders(sec, chunk_size=chunk_size))
                        else:
                            placeholder_chunks.append(sec)
                        for p in sec:
                            if p in seen:
                                continue
                            seen.add(p)
                            required_placeholders.append(p)
                else:
                    required_placeholders = _extract_required_placeholders(app_config.templates.proposal)
        if not placeholder_chunks:
            base_placeholders = required_placeholders or list(PLACEHOLDER_FIELDS)
            placeholder_chunks = _chunk_placeholders(base_placeholders, chunk_size=chunk_size)

        stage = "prepare_tables"
        gen_table_chunks: list[list[str]] = []
        seen_gen_tables: set[str] = set()
        for chunk in table_chunks:
            if isinstance(chunk, list) and chunk:
                gen_table_chunks.append(chunk)
                for name in chunk:
                    if isinstance(name, str) and name.strip():
                        seen_gen_tables.add(name)
        for name in TABLE_MIN_SPECS.keys():
            if name not in seen_gen_tables:
                gen_table_chunks.append([name])

        stage = "build_graph"
        graph = build_graph(
            ledger_runtime=ledger_runtime,
            final_runtime=final_runtime,
            section_chunks=placeholder_chunks,
            table_chunks=gen_table_chunks,
        )
        compiled = graph.compile()
        # Required tables = template table order + TABLE_MIN_SPECS keys (base set)
        table_order: list[str] = []
        seen_tables: set[str] = set()
        for chunk in table_chunks:
            if not isinstance(chunk, list):
                continue
            for name in chunk:
                if not isinstance(name, str) or not name.strip():
                    continue
                if name in seen_tables:
                    continue
                seen_tables.add(name)
                table_order.append(name)
        for name in TABLE_MIN_SPECS.keys():
            if name not in seen_tables:
                table_order.append(name)
        required_tables = table_order
        state = {
            "spec_text": spec_text,
            "manual_inputs": manual_inputs,
            "required_placeholders": required_placeholders,
            "required_tables": required_tables,
            "section_outputs": [],
            "locked_placeholders": {},
            "locked_tables": {},
        }

        stage = "invoke_graph"
        result = compiled.invoke(state)
        llm_output = result.get("llm_output", {}) if isinstance(result, dict) else {}
        ledger = result.get("ledger", {}) if isinstance(result, dict) else {}
        metrics = result.get("metrics", {}) if isinstance(result, dict) else {}

        stage = "build_placeholder_map"
        placeholder_map = build_placeholder_map(manual_inputs, llm_output)

        stage = "metrics_output"
        out_dir = args.out
        runs_env = os.getenv("PROPOSAL_RUNS_LOG", "").strip().lower()
        runs_flag = getattr(args, "runs_log", False) or runs_env in {"1", "true", "yes", "y"}
        if runs_flag:
            _write_metrics_log(metrics or {}, out_dir)
        _print_metrics_summary(metrics or {})

        stage = "render_docx"
        out_path = _render_docx(app_config, manual_inputs, llm_output, placeholder_map, out_dir)
        return {
            "out_path": out_path,
            "llm_output": llm_output,
            "placeholder_map": placeholder_map,
        }
    except Exception as exc:
        error_info = {
            "type": type(exc).__name__,
            "message": str(exc),
            "stage": stage,
        }
        if not isinstance(metrics, dict):
            metrics = {}
        metrics.setdefault("error", error_info)
        raise
    finally:
        if debug_flag:
            safe_llm_output = llm_output if isinstance(llm_output, dict) else {}
            safe_placeholder_map = placeholder_map if isinstance(placeholder_map, dict) else {}
            safe_ledger = ledger if isinstance(ledger, dict) else {}
            safe_metrics = metrics if isinstance(metrics, dict) else {}
            if error_info:
                safe_metrics.setdefault("error", error_info)
            _write_debug(debug_dir, safe_llm_output, safe_placeholder_map, safe_ledger, safe_metrics)
