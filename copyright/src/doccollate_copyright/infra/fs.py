from __future__ import annotations

import logging
import os
import subprocess
import tempfile
from pathlib import Path

import yaml
from docx import Document

from .path_utils import normalize_path, normalize_path_string

logger = logging.getLogger(__name__)


def read_file_content(file_path: Path, max_chars: int = 40000) -> str:
    file_path = normalize_path(file_path)
    ext = file_path.suffix.lower()
    text = ""
    try:
        if ext == ".md":
            text = file_path.read_text(encoding="utf-8")[:max_chars]
        elif ext == ".docx":
            text = _read_docx_via_pandoc(file_path, max_chars)
            if not text:
                text = _read_docx_fallback(file_path, max_chars)
        elif ext == ".pdf":
            text = _read_pdf_text(file_path, max_chars)
    except Exception as exc:
        logger.error("Failed to read %s: %s", file_path, exc)
        return ""
    return text


def _read_docx_via_pandoc(file_path: Path, max_chars: int) -> str:
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            md_path = Path(tmp_dir) / f"{file_path.stem}.md"
            result = subprocess.run(
                ["pandoc", str(file_path), "-t", "gfm", "-o", str(md_path)],
                capture_output=True,
                text=True,
            )
            if result.returncode != 0:
                stderr = result.stderr.strip() or result.stdout.strip()
                logger.warning("Pandoc conversion failed for %s: %s", file_path, stderr)
                return ""
            content = md_path.read_text(encoding="utf-8")
            save_path = _pandoc_save_path(file_path)
            if save_path:
                try:
                    save_path.parent.mkdir(parents=True, exist_ok=True)
                    save_path.write_text(content, encoding="utf-8")
                except Exception as exc:
                    logger.warning("Failed to save converted md for %s: %s", file_path, exc)
            return content[:max_chars]
    except FileNotFoundError:
        logger.warning("Pandoc not found in PATH. Falling back to basic docx parsing.")
    except Exception as exc:
        logger.warning("Pandoc conversion failed for %s: %s", file_path, exc)
    return ""


def _read_docx_fallback(file_path: Path, max_chars: int) -> str:
    doc = Document(file_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        content = paragraph.text.strip()
        if content:
            parts.append(content)
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                row_cells.append(cell_text)
            if any(row_cells):
                parts.append(" | ".join(cell for cell in row_cells if cell))
    return "\n".join(parts)[:max_chars]


def _read_pdf_text(file_path: Path, max_chars: int) -> str:
    try:
        import pdfplumber
    except Exception:  # pragma: no cover
        logger.warning("Missing dependency: pdfplumber")
        return ""
    parts: list[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text:
                    parts.append(text)
    except Exception as exc:
        logger.warning("PDF text extraction failed for %s: %s", file_path, exc)
    return "\n".join(parts)[:max_chars]


def _pandoc_save_path(file_path: Path) -> Path | None:
    flag = os.getenv("DOCCOLLATE_PANDOC_SAVE_MD", "").strip().lower()
    if flag not in {"1", "true", "yes", "y"}:
        return None
    output_dir = os.getenv("DOCCOLLATE_PANDOC_SAVE_DIR", "").strip()
    if output_dir:
        return (Path(output_dir).expanduser() / f"{file_path.stem}.md").resolve()
    return file_path.with_suffix(".md")


def load_yaml_config(path: Path) -> dict:
    path = normalize_path(path)
    if not path.exists():
        logger.warning("Config file not found: %s", path)
        return {}
    try:
        with path.open("r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception as exc:
        logger.error("Failed to load config: %s", exc)
        return {}
