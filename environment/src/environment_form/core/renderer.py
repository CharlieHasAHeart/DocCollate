from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


def _replace(doc: Document, replacements: dict[str, str]) -> None:
    def replace_in_paragraph(p):
        while "{{" in p.text:
            old = p.text
            for k, v in replacements.items():
                p.text = p.text.replace(k, v)
                p.text = p.text.replace(k.replace("}}", "}}\\v"), v)
            if old == p.text:
                break

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)


def generate_document(template_path: Path, output_path: Path, data: dict[str, str]) -> None:
    try:
        doc = Document(template_path)
    except Exception as exc:
        logger.error("Unable to open template '%s': %s", template_path, exc)
        raise

    replacements = {
        "{{env__server_os}}": data.get("env__server_os", ""),
        "{{env__server_soft}}": data.get("env__server_soft", ""),
        "{{env__server_model}}": data.get("env__server_model", ""),
        "{{env__server_config}}": data.get("env__server_config", ""),
        "{{env__server_id}}": data.get("env__server_id", ""),
        "{{env__client_os}}": data.get("env__client_os", ""),
        "{{env__client_soft}}": data.get("env__client_soft", ""),
        "{{env__client_model}}": data.get("env__client_model", ""),
        "{{env__client_config}}": data.get("env__client_config", ""),
        "{{env__client_id}}": data.get("env__client_id", ""),
    }
    _replace(doc, replacements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
